#!/usr/bin/env python3
"""
教学文档 Word 解析工具
- 读取Word文档，提取结构化内容
- 生成/填入Word文档
- 随机生成AB卷
"""

import os
import json
import sys
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from docx import Document
except ImportError:
    print("需要安装 python-docx: pip install python-docx")
    sys.exit(1)


class TeachingDocParser:
    """教学文档解析器"""

    # 关键词映射表
    FIELD_PATTERNS = {
        'course_name': ['课程名称', '课程名称：'],
        'course_code': ['课程编号'],
        'credit': ['学分'],
        'hours': ['学时', '总学时'],
        'semester': ['授课学期'],
        'department': ['学院', '系部'],
        'major': ['专业'],
        'teacher': ['授课教师', '主讲教师', '制定教师'],
        'class_info': ['授课对象', '授课班级'],
        'course_type': ['课程类别', '课程类型'],
        'course_intro': ['课程简介', '课程内容', '（一）课程内容'],
        'design理念': ['课程设计理念'],
        'quality_goal': ['素质目标'],
        'knowledge_goal': ['知识目标'],
        'ability_goal': ['能力目标'],
        'assessment': ['考核方式', '考核内容'],
        'references': ['参考书', '推荐教材'],
        'teaching_content': ['教学内容', '教学内容组织'],
        'teaching_method': ['教学方法'],
        'teaching_plan': ['授课计划', '教学进度'],
        'teaching_process': ['教学过程', '教学过程设计'],
        'teaching_resource': ['学习资源', '课程资源'],
        'key_point': ['教学重点', '重点'],
        'difficult_point': ['教学难点', '难点'],
    }

    def __init__(self, file_path):
        self.file_path = file_path
        self.doc = None
        self.paragraphs = []
        self.tables = []

    def load(self):
        """加载Word文档"""
        self.doc = Document(self.file_path)
        self.paragraphs = [p.text.strip() for p in self.doc.paragraphs if p.text.strip()]
        self.tables = self.doc.tables
        return self

    def parse_outline(self):
        """解析教学大纲"""
        content = {}
        current_section = None
        current_content = []

        for para in self.paragraphs:
            matched = False
            for field, patterns in self.FIELD_PATTERNS.items():
                for pattern in patterns:
                    if pattern in para and len(para) < 50:
                        if current_section:
                            content[current_section] = '\n'.join(current_content).strip()
                        current_section = field
                        current_content = []
                        matched = True
                        break
                if matched:
                    break

            if not matched and current_section:
                if len(para) > 5 and not self._is_noise(para):
                    current_content.append(para)

        if current_section:
            content[current_section] = '\n'.join(current_content).strip()

        for table in self.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if rows:
                content['table_data'] = rows

        return content

    def parse_exam(self):
        """解析试卷，提取题目"""
        questions = []
        current_type = None
        current_content = []

        type_keywords = {
            '选择': r'[一二三四1-4]、?选择题',
            '填空': r'[一二三四1-4]、?填空题',
            '简答': r'[一二三四1-4]、?简答题',
            '论述': r'[一二三四1-4]、?论述题',
            '实践': r'[一二三四1-4]、?实践题',
        }

        for para in self.paragraphs:
            for qtype, pattern in type_keywords.items():
                if re.search(pattern, para):
                    if current_type and current_content:
                        questions.append({
                            'type': current_type,
                            'content': '\n'.join(current_content)
                        })
                    current_type = qtype
                    current_content = []
                    break

            if current_type and not any(k in para for k in type_keywords.keys()):
                if re.match(r'^\d+[.、)]', para) or re.match(r'^[（\(]\d+[）\)]', para):
                    if current_content:
                        questions.append({
                            'type': current_type,
                            'content': '\n'.join(current_content)
                        })
                        current_content = []
                    current_content.append(re.sub(r'^[（\(]?\d+[.、)）\)]', '', para))
                elif current_content:
                    current_content.append(para)

        if current_type and current_content:
            questions.append({
                'type': current_type,
                'content': '\n'.join(current_content)
            })

        return questions

    def _is_noise(self, text):
        """判断是否是噪音行"""
        noise_patterns = [
            r'^注：', r'^备注：', r'^说明：', r'^请[按各]', r'^此表',
        ]
        for pattern in noise_patterns:
            if re.search(pattern, text):
                return True
        return False


class WordGenerator:
    """Word文档生成器"""

    def __init__(self):
        self.doc = Document()

    def add_title(self, text, level=1):
        heading = self.doc.add_heading(text, level=level)
        return heading

    def add_paragraph(self, text, bold=False, size=12):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        return p

    def add_field(self, label, value):
        p = self.doc.add_paragraph()
        run1 = p.add_run(f'{label}：')
        run1.bold = True
        run1.font.size = Pt(12)
        run2 = p.add_run(str(value) if value else '')
        run2.font.size = Pt(12)
        return p

    def generate_exam_paper(self, title, questions, total_score=100):
        """生成试卷"""
        self.add_title(title, 1)
        self.add_paragraph(f'总分：{total_score}分')
        self.doc.add_paragraph()

        type_names = {
            '选择': '一、选择题',
            '填空': '二、填空题',
            '简答': '三、简答题',
            '论述': '四、论述题',
            '实践': '五、实践题'
        }

        for qtype, items in questions.items():
            if items:
                self.add_title(type_names.get(qtype, qtype), 2)
                for i, q in enumerate(items, 1):
                    self.add_paragraph(f'{i}. {q.get("content", "")}')

        return self.doc

    def save(self, output_path):
        self.doc.save(output_path)
        print(f'已保存到: {output_path}')


def main():
    if len(sys.argv) < 3:
        print('用法:')
        print('  解析大纲: python doc_parser.py parse outline <input.docx> [output.json]')
        print('  解析试卷: python doc_parser.py parse exam <input.docx> [output.json]')
        print('  生成试卷: python doc_parser.py generate <config.json> <output.docx>')
        sys.exit(1)

    command = sys.argv[1]

    if command == 'parse':
        action = sys.argv[2]
        input_file = sys.argv[3]
        output_file = sys.argv[4] if len(sys.argv) > 4 else None

        parser = TeachingDocParser(input_file)
        parser.load()

        if action == 'outline':
            result = parser.parse_outline()
        elif action == 'exam':
            result = parser.parse_exam()
        else:
            print(f'未知的解析类型: {action}')
            sys.exit(1)

        if output_file:
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(result, f, ensure_ascii=False, indent=2)
            print(f'结果已保存到: {output_file}')
        else:
            print(json.dumps(result, ensure_ascii=False, indent=2))

    elif command == 'generate':
        config_file = sys.argv[2]
        output_file = sys.argv[3]

        with open(config_file, 'r', encoding='utf-8') as f:
            config = json.load(f)

        gen = WordGenerator()
        doc = gen.generate_exam_paper(
            config['title'],
            config['questions'],
            config.get('total_score', 100)
        )
        gen.save(output_file)


if __name__ == '__main__':
    main()
